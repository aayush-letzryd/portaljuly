import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Define path to Downloads folder
user_home = os.path.expanduser("~")
downloads_folder = os.path.join(user_home, "Downloads")
output_file = os.path.join(downloads_folder, "LetzRyd_Portal_Primary_Credentials.docx")

# Create Document
doc = docx.Document()

# Set standard page margins (0.8 inch)
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Helper function to set cell padding
def set_cell_padding(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

# Helper function to set cell background color
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

# Helper function to set table borders / gridlines
def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

# Title
p_title = doc.add_paragraph()
p_title.paragraph_format.space_after = Pt(2)
run_title = p_title.add_run("LetzRyd Web Portal Credentials")
run_title.font.name = "Inter"
run_title.font.size = Pt(18)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(15, 23, 42)

# Subtitle / Note
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(14)
run_sub = p_sub.add_run("Default password for all accounts: 123456")
run_sub.font.name = "Inter"
run_sub.font.size = Pt(11)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(51, 65, 85)

# 15 Primary Data Rows
table_data = [
    [3, "admin", "SA", "Super Admin", "System Super Admin", "123456"],
    [41, "general_manager", "GM", "General Manager", "General Manager 1", "123456"],
    [17, "business", "BH", "Business Head", "Business Head 1", "123456"],
    [18, "finance_lead", "FL", "Finance Lead", "Finance Lead 1", "123456"],
    [19, "finance_executive", "FE", "Finance Executive", "Finance Executive 1", "123456"],
    [20, "city_manager", "CM", "City Manager", "City Manager 1", "123456"],
    [21, "driver_manager", "DM", "Driver Manager", "Driver Manager 1", "123456"],
    [23, "ops_executive", "OE", "Ops Executive", "Ops Executive 1", "123456"],
    [24, "fleet_manager", "FM", "Fleet Manager", "Fleet Manager 1", "123456"],
    [25, "maintenance_coordinator", "MC", "Maintenance Coordinator", "Maintenance Coordinator 1", "123456"],
    [26, "onboarding_executive", "OB", "Onboarding Executive", "Onboarding Executive 1", "123456"],
    [28, "support_executive", "SP", "Support Executive", "Support Executive 1", "123456"],
    [29, "auditor", "AU", "Auditor / Compliance", "Auditor 1", "123456"],
    [31, "fleet_partner", "PT", "Fleet Partner / Operator", "Fleet Partner 1", "123456"],
    [32, "driver", "DR", "Driver", "Driver 1", "123456"]
]

headers = ["User ID", "Username", "Role Code", "Role Name", "Generic Employee Title", "Password"]
col_widths_in = [0.8, 1.8, 0.9, 1.8, 2.0, 1.0]

main_table = doc.add_table(rows=len(table_data) + 1, cols=6)
main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(main_table, color="CBD5E1", sz="4", val="single") # Crisp slate gridlines

# Format Header Row (Subtle neutral grey header, clear bold text)
hdr_cells = main_table.rows[0].cells
for idx, header_text in enumerate(headers):
    hdr_cells[idx].text = header_text
    set_cell_background(hdr_cells[idx], "F1F5F9") # Subtle light grey
    set_cell_padding(hdr_cells[idx], top=140, bottom=140, left=120, right=120)
    p = hdr_cells[idx].paragraphs[0]
    if idx in [0, 2, 5]:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)

# Populate Data Rows (Clean crisp grid cells, Inter font)
for r_idx, r_data in enumerate(table_data, start=1):
    row_cells = main_table.rows[r_idx].cells
    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
    
    for c_idx, val in enumerate(r_data):
        cell = row_cells[c_idx]
        cell.text = str(val)
        set_cell_background(cell, bg)
        set_cell_padding(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        
        if c_idx in [0, 2, 5]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        for run in p.runs:
            run.font.name = "Inter"
            run.font.size = Pt(9.5)
            if c_idx == 1: # Username
                run.font.bold = True
                run.font.color.rgb = RGBColor(15, 23, 42)
            else:
                run.font.color.rgb = RGBColor(51, 65, 85)

# Apply column widths
for row in main_table.rows:
    for idx, width in enumerate(col_widths_in):
        row.cells[idx].width = Inches(width)

# Save Word document
doc.save(output_file)
print(f"[OK] Structured Gridlines Inter Word Document (.docx) saved to: {output_file}")
