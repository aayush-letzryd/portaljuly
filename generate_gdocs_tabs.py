import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Define target folder in Downloads
user_home = os.path.expanduser("~")
downloads_folder = os.path.join(user_home, "Downloads")
tabs_dir = os.path.join(downloads_folder, "LetzRyd_Portal_Google_Docs_Tabs")

os.makedirs(tabs_dir, exist_ok=True)

def create_tab_doc(file_name, tab_title, sections_content):
    doc = docx.Document()
    
    # Margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    # Title
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_after = Pt(2)
    r_t = p_t.add_run(tab_title)
    r_t.font.name = "Inter"
    r_t.font.size = Pt(18)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(15, 23, 42)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("LetzRyd Web Portal — Google Docs Modular Documentation Tab")
    r_sub.font.name = "Inter"
    r_sub.font.size = Pt(10)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    for sec_title, text_blocks, table_info in sections_content:
        # Heading
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(14)
        p_h.paragraph_format.space_after = Pt(4)
        p_h.paragraph_format.keep_with_next = True
        r_h = p_h.add_run(sec_title)
        r_h.font.name = "Inter"
        r_h.font.size = Pt(13)
        r_h.font.bold = True
        r_h.font.color.rgb = RGBColor(30, 41, 59)
        
        # Text blocks
        for b in text_blocks:
            p_b = doc.add_paragraph()
            p_b.paragraph_format.space_after = Pt(4)
            p_b.paragraph_format.line_spacing = 1.15
            r_b = p_b.add_run(b)
            r_b.font.name = "Inter"
            r_b.font.size = Pt(9.5)
            r_b.font.color.rgb = RGBColor(51, 65, 85)
            
        # Optional table
        if table_info:
            headers, rows_data, col_w = table_info
            table = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Borders
            tblPr = table._element.xpath('w:tblPr')
            if tblPr:
                borders = parse_xml(f'''
                    <w:tblBorders {nsdecls("w")}>
                        <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:left w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                        <w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                    </w:tblBorders>
                ''')
                tblPr[0].append(borders)
                
            # Headers
            for c_idx, h_text in enumerate(headers):
                cell = table.rows[0].cells[c_idx]
                cell.text = h_text
                tcPr = cell._element.get_or_add_tcPr()
                tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>'))
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Inter"
                    r.font.size = Pt(9.5)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(15, 23, 42)
                    
            # Data rows
            for r_idx, r_val in enumerate(rows_data, start=1):
                row_cells = table.rows[r_idx].cells
                bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c_idx, val in enumerate(r_val):
                    c_cell = row_cells[c_idx]
                    c_cell.text = str(val)
                    tcPr = c_cell._element.get_or_add_tcPr()
                    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'))
                    p = c_cell.paragraphs[0]
                    if c_idx in [0, 2]:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.name = "Inter"
                        r.font.size = Pt(9)
                        if c_idx == 1:
                            r.font.bold = True
                            
            for r in table.rows:
                for idx, w in enumerate(col_w):
                    r.cells[idx].width = Inches(w)
                    
    target_path = os.path.join(tabs_dir, file_name)
    doc.save(target_path)
    print(f"[OK] Created Tab Doc: {target_path}")

# TAB 1: ARCHITECTURE & STACK
t1_content = [
    ("1. High-Level Architecture Overview", [
        "The LetzRyd Web Portal is engineered as an enterprise-grade Single Page Application (SPA) backed by a Python FastAPI asynchronous REST API service and a PostgreSQL relational database.",
        "Architecture Components: Frontend SPA (React 18 + Vite + TypeScript), Backend REST API (Python 3.12 + FastAPI + Uvicorn), Database Tier (PostgreSQL with psycopg2 ThreadedConnectionPool)."
    ], None),
    ("2. Technology Stack Breakdown", [
        "Detailed framework and runtime specifications:"
    ], (
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Frontend Core", "React JS", "18.3.1", "UI component rendering and state management"],
            ["Frontend Build", "Vite JS", "6.4.3", "HMR development server & production builder"],
            ["Type Safety", "TypeScript", "5.5.3", "Strict types & interface definitions"],
            ["Backend API", "FastAPI / Python", "0.109.0 / 3.12", "Asynchronous REST endpoints & validation"],
            ["Database", "PostgreSQL", "15.0+", "Relational storage & JSONB audit logs"]
        ],
        [1.5, 1.8, 1.0, 2.5]
    ))
]
create_tab_doc("Tab_1_Architecture_and_Stack.docx", "Tab 1 — Architecture & Technology Stack", t1_content)

# TAB 2: FORMS & SPECIFICATIONS
t2_content = [
    ("1. Partner Onboarding Form & KYC", [
        "Captures candidate personal info, mandatory live selfie, Aadhaar front/back photos, PAN photo, PAN-Aadhaar link toggle, 1-4 page address proof uploads, bank passbook image, emergency contacts, and ride-hailing experience."
    ], None),
    ("2. Vehicle Allocation Form", [
        "Transaction Types: New Allocation, Reallocation, Rejoining, Swap.",
        "Auto-fetch driver details by 10-digit mobile lookup, vehicle autocomplete, GPS status toggle, OLA negative balance + proof photo, odometer reading + photo, 5 condition photos (Left, Right, Front, Back, Battery), accessories inspection checklist."
    ], None),
    ("3. Vehicle Drop-Off Form", [
        "Return Reasons: Voluntary Return, Default, Breakdown, Contract Completion.",
        "Odometer + 5 condition photos, pending dues, damage penalty, deposit refund status (Pending Assessment, Approved, Deductions, Forfeited), Saved Drafts tab, live search, city/reason filters, 10/page pagination, CSV export."
    ], None),
    ("4. Walk-In & Returning Partner Visit Logger", [
        "Two-Table Architecture: july_new_walkins (candidates) and july_existing_walkins (returning partners).",
        "Lookup by phone number auto-fills partner name and city, tags operational visit reason, and displays prior visit history modal."
    ], None)
]
create_tab_doc("Tab_2_Forms_and_Specifications.docx", "Tab 2 — Form Modules & Specifications", t2_content)

# TAB 3: RBAC & PERMISSIONS MATRIX
t3_content = [
    ("1. System Roles & Access Control", [
        "Granular access control enforcing permissions across all 15 system roles."
    ], (
        ["User ID", "Username", "Role Code", "Role Name", "Title", "Primary Form Access"],
        [
            [3, "admin", "SA", "Super Admin", "System Super Admin", "Full Access to All Forms & Settings"],
            [41, "general_manager", "GM", "General Manager", "General Manager 1", "Global Registries & High-Level Approvals"],
            [17, "business", "BH", "Business Head", "Business Head 1", "Business Overview & Approvals Dashboard"],
            [18, "finance_lead", "FL", "Finance Lead", "Finance Lead 1", "Financial Approvals & Payouts"],
            [19, "finance_executive", "FE", "Finance Executive", "Finance Executive 1", "Dues Assessment & Verification"],
            [20, "city_manager", "CM", "City Manager", "City Manager 1", "City Approvals & Allocation Registries"],
            [21, "driver_manager", "DM", "Driver Manager", "Driver Manager 1", "Driver Onboarding & Walk-In Logger"],
            [23, "ops_executive", "OE", "Ops Executive", "Ops Executive 1", "Vehicle Allocation & Drop-Off Forms"],
            [24, "fleet_manager", "FM", "Fleet Manager", "Fleet Manager 1", "Fleet Master Logs & Allocations"],
            [25, "maintenance_coordinator", "MC", "Maintenance Coordinator", "Maintenance Coordinator 1", "Vehicle Repairs & Drop-Off Inspection"],
            [26, "onboarding_executive", "OB", "Onboarding Executive", "Onboarding Executive 1", "Onboarding Form & Candidate Walk-Ins"],
            [28, "support_executive", "SP", "Support Executive", "Support Executive 1", "Walk-In Logger & Support Tickets"],
            [29, "auditor", "AU", "Auditor / Compliance", "Auditor 1", "Read-Only View Access"],
            [31, "fleet_partner", "PT", "Fleet Partner", "Fleet Partner 1", "Operator Fleet View"],
            [32, "driver", "DR", "Driver", "Driver 1", "Personal History & Profile View"]
        ],
        [0.7, 1.4, 0.8, 1.4, 1.4, 1.4]
    ))
]
create_tab_doc("Tab_3_RBAC_and_User_Permissions.docx", "Tab 3 — RBAC & Permissions Matrix", t3_content)

# TAB 4: AUDIT LOGS
t4_content = [
    ("1. Audit Log Tables Specification", [
        "The system incorporates 6 active log tables capturing real-time field-level JSON diffs and IST timestamps."
    ], (
        ["Audit Log Table", "Module Tracked", "Actions Logged", "Data Recorded"],
        [
            ["july_allocation_form_logs", "Allocation & Drop-Off", "CREATE, UPDATE, DELETE", "Allocation ID, Executive ID, IST Timestamp, JSON diff"],
            ["july_walkin_logs", "Walk-In & Visits", "CREATE, UPDATE, DELETE", "Walkin ID (N1, E1), Executive ID, IST Timestamp, Remarks"],
            ["july_onboarding_logs", "Partner Onboarding", "CREATE, UPDATE, DRAFT", "Onboarding ID, Executive ID, Full Snapshot & JSON diff"],
            ["july_vehicle_logs", "Fleet Master", "CREATE, UPDATE, STATUS", "Vehicle ID, Status, Executive ID, Snapshot & JSON diff"],
            ["july_approval_chain_logs", "Approvals Workflow", "APPROVE, REJECT, REVISION", "Module, Record ID, From/To User, Remarks & Timestamp"],
            ["july_user_login_logs", "Authentication Engine", "LOGIN, LOGOUT", "User ID, Username, Role Code, IP Address & Session Token"]
        ],
        [1.8, 1.5, 1.5, 2.0]
    ))
]
create_tab_doc("Tab_4_Audit_Logs_and_Tracking.docx", "Tab 4 — Audit Logs & Event Engine", t4_content)

# TAB 5: DATABASE SCHEMA & WORKFLOWS
t5_content = [
    ("1. Database Schema & Tables Overview", [
        "july_portal_users: Executive accounts, roles, employee IDs, cities, joining dates.",
        "july_form_onboarding: Master candidate onboarding registry & KYC document URLs.",
        "july_allocation_form: Vehicle allocation and vehicle drop-off entries.",
        "july_new_walkins & july_existing_walkins: Walk-in candidate and returning partner visit logs."
    ], None),
    ("2. Approval Workflows State Machine", [
        "Draft State -> Pending Approval State -> Approved State / Rejected State / Revision Required State."
    ], None)
]
create_tab_doc("Tab_5_Database_Schema_and_Workflows.docx", "Tab 5 — Database Schema & Workflows", t5_content)

# TAB 6: DEPLOYMENT & DEVOPS
t6_content = [
    ("1. Local Development Commands", [
        "Backend Uvicorn Server: .venv\\Scripts\\python.exe -m uvicorn main:app --reload --host 127.0.0.1 --port 8000",
        "Frontend Vite Server: npm run dev (http://localhost:3000)",
        "Production Build: npm run build"
    ], None),
    ("2. Environment Variables & Database Connection Pooling", [
        "DATABASE_URL=postgresql://postgres:postgres@localhost:5432/portal_db",
        "PostgreSQL ThreadedConnectionPool sizing: min 1, max 20 connections."
    ], None)
]
create_tab_doc("Tab_6_Deployment_and_DevOps.docx", "Tab 6 — Deployment & DevOps Guide", t6_content)

print("[OK] All 6 Google Docs Modular Tab files generated successfully in:", tabs_dir)
