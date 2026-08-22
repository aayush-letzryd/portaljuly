import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Define path to Downloads folder
user_home = os.path.expanduser("~")
downloads_folder = os.path.join(user_home, "Downloads")
output_file = os.path.join(downloads_folder, "LetzRyd_Web_Portal_Comprehensive_Documentation.docx")

# Create Document
doc = docx.Document()

# Set standard page margins (0.8 inch)
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Helper function to set cell padding
def set_cell_padding(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

# Helper function to set cell background color
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

# Helper function to set table borders / gridlines
def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)
    return p

def add_paragraph(text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Inter"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r_bullet = p.add_run("• ")
    r_bullet.font.name = "Inter"
    r_bullet.font.size = Pt(10)
    r_bullet.font.bold = True
    r_bullet.font.color.rgb = RGBColor(15, 23, 42)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix + ": ")
        r_pre.font.name = "Inter"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
        
    run = p.add_run(text)
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

# DOCUMENT HEADER
p_title = doc.add_paragraph()
p_title.paragraph_format.space_after = Pt(2)
run_title = p_title.add_run("LetzRyd Web Portal — System Architecture & Technical Specifications")
run_title.font.name = "Inter"
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(15, 23, 42)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(16)
run_sub = p_sub.add_run("Comprehensive Technical Documentation for Forms, RBAC, Audit Logs, Database Schema, Workflows & Deployment")
run_sub.font.name = "Inter"
run_sub.font.size = Pt(11)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(71, 85, 105)

# SECTION 1: EXECUTIVE SUMMARY & SYSTEM ARCHITECTURE
add_heading_1("1. Executive Summary & System Architecture")
add_paragraph("The LetzRyd Web Portal is an enterprise-grade operational management web application designed for fleet operations, driver onboarding, vehicle allocation, vehicle drop-offs, candidate lead tracking, and multi-tier approval workflows across LetzRyd operating hubs.")

add_heading_2("1.1 High-Level Architecture Overview")
add_paragraph("The platform is engineered as a high-performance Single Page Application (SPA) backed by a Python FastAPI asynchronous backend REST service and PostgreSQL relational database.")

add_bullet("Frontend Architecture", "Built with React 18, Vite 6, and TypeScript. Styled using Vanilla CSS tokens and Tailwind CSS utilities, utilizing Lucide React icons for consistent visual cues.")
add_bullet("Backend Microservices", "Powered by Python 3.12 and FastAPI. Implements non-blocking REST API endpoints, Pydantic data contract validation, and connection pooling via psycopg2.")
add_bullet("Database & Storage", "PostgreSQL database utilizing connection pooling (min 1, max 20 connections), JSONB column types for unstructured photo metadata & audit diffs, GIN indexing, and IST timezone synchronization.")
add_bullet("Security & RBAC", "Session-based token authentication backed by july_app_sessions table, granular 15-role permission matrix, and parameterized SQL queries to prevent SQL injection.")

# SECTION 2: TECHNOLOGY STACK & DEPENDENCIES
add_heading_1("2. Technology Stack & Environment Configuration")
add_paragraph("Below is the complete breakdown of software frameworks, libraries, runtime environments, and database configurations utilized in the portal.")

tech_headers = ["Layer", "Technology / Package", "Version", "Purpose & Utility"]
tech_data = [
    ["Frontend Core", "React JS", "18.3.1", "UI component tree rendering, state management, and DOM handling"],
    ["Frontend Build", "Vite JS", "6.4.3", "Ultra-fast HMR development server and production bundle builder"],
    ["Type Safety", "TypeScript", "5.5.3", "Strict type checking, interfaces, and component prop definitions"],
    ["Styling Engine", "Tailwind CSS + CSS", "3.4.1", "Utility classes, glassmorphism overlays, and modern visual design"],
    ["Iconography", "Lucide React", "0.344.0", "UI icons for navigation, actions, status badges, and controls"],
    ["Backend REST API", "FastAPI / Python", "0.109.0 / 3.12", "Asynchronous API routing, request validation, and middleware"],
    ["ASGI Web Server", "Uvicorn", "0.27.0", "ASGI web server running main FastAPI application on port 8000"],
    ["Data Validation", "Pydantic", "2.6.1", "Request payload schema validation, serialization, and type parsing"],
    ["Database Driver", "psycopg2-binary", "2.9.9", "PostgreSQL database driver with ThreadedConnectionPool"],
    ["Database Engine", "PostgreSQL", "15.0+", "Relational database storing form submissions, logs, users & state"],
    ["Password Hashing", "Passlib (Bcrypt)", "1.7.4", "Secure password hashing and flexible verification mechanisms"]
]

t_tech = doc.add_table(rows=len(tech_data)+1, cols=4)
t_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_tech)

hdr_c = t_tech.rows[0].cells
for idx, text in enumerate(tech_headers):
    hdr_c[idx].text = text
    set_cell_background(hdr_c[idx], "F1F5F9")
    set_cell_padding(hdr_c[idx])
    p = hdr_c[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Inter"
        r.font.size = Pt(9.5)
        r.font.bold = True

for r_i, r_d in enumerate(tech_data, start=1):
    row_c = t_tech.rows[r_i].cells
    bg = "F8FAFC" if r_i % 2 == 1 else "FFFFFF"
    for c_i, val in enumerate(r_d):
        row_c[c_i].text = str(val)
        set_cell_background(row_c[c_i], bg)
        set_cell_padding(row_c[c_i])
        p = row_c[c_i].paragraphs[0]
        if c_i in [0, 2]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Inter"
            r.font.size = Pt(9)
            if c_i == 1:
                r.font.bold = True

for r in t_tech.rows:
    r.cells[0].width = Inches(1.3)
    r.cells[1].width = Inches(1.8)
    r.cells[2].width = Inches(1.0)
    r.cells[3].width = Inches(3.4)

# SECTION 3: CORE MODULES & FORM SPECIFICATIONS
add_heading_1("3. Core Modules & Detailed Form Specifications")
add_paragraph("The LetzRyd portal comprises four core operational form modules designed to capture candidate onboarding, vehicle allocation, vehicle drop-offs, and hub walk-in visits.")

add_heading_2("3.1 Partner Onboarding Form & KYC Module")
add_paragraph("Captures complete personal details, identity verification documents, multi-page address proofs, banking credentials, emergency contacts, and prior platform experience for new drivers and fleet partners.")

add_bullet("Candidate Basic Info", "First Name, Last Name, Full Name, Mobile Number (10-digit validation), Operating City, Department.")
add_bullet("Driver Selfie Capture", "Mandatory live camera capture or image file upload for driver face verification during onboarding.")
add_bullet("Aadhaar Identity Verification", "Separate upload inputs for Aadhaar Card Front Photo and Aadhaar Card Back Photo with image preview.")
add_bullet("PAN Card & Verification Toggle", "PAN Card Photo upload slot and checkbox toggle verifying 'Is Aadhaar linked with PAN'.")
add_bullet("Multi-Page Local Address Proof", "Upload slots for 1 to 4 pages of local address proof (Rent agreement, electricity bill, gas bill).")
add_bullet("Bank Details & Passbook Image", "Bank Account Number, IFSC Code, Account Holder Name, and Cancelled Cheque / Bank Passbook photo upload.")
add_bullet("Emergency Contacts", "Contact Person Name, 10-digit Mobile Number, and Relationship dropdown (Father, Spouse, Brother, Friend).")
add_bullet("Third-Party Platform Dropdown", "Single-select dropdown for prior ride-hailing experience: Uber, Ola, Rapido, or None.")

add_heading_2("3.2 Vehicle Allocation Form Module")
add_paragraph("Handles vehicle assignment logistics, meter readings, car condition photos, inspection checklists, and vehicle swap returns.")

add_bullet("Transaction Type Selector", "Dropdown choices: New Allocation, Reallocation, Rejoining, and Swap.")
add_bullet("Driver Phone & ID Auto-Fetch", "Driver Fetch button by phone or ID that auto-populates driver name, city, rental plan, contract type, and car model.")
add_bullet("Vehicle Autocomplete", "Auto-capitalizing vehicle registration number input with live suggestions from active fleet database.")
add_bullet("GPS Active Status", "Dropdown verification status (Yes / No) confirming GPS tracker device is operational before release.")
add_bullet("OLA Balance & Meter Capture", "OLA Negative Balance amount field (₹), OLA Balance Proof Photo upload, Odometer Reading (KM), and mandatory Odometer Photo upload.")
add_bullet("Car Condition Photos (5 Slots)", "5 mandatory photo upload slots: Left Side, Right Side, Front Side, Back Side, and Battery Photo.")
add_bullet("Accessories Inspection Checklist", "Interactive checklist for Jack, Jack Rod, Spanner, Parking Triangle, Fire Extinguisher, Seat Cover, Floor Carpet, and Music System.")
add_bullet("Swap Vehicle Return Details", "For vehicle swaps: captures Old Vehicle Number, Drop-Off Odometer, Return Remarks, and Returned Vehicle Inspection Checklist.")

add_heading_2("3.3 Vehicle Drop-Off Form Module")
add_paragraph("Records vehicle returns at hubs or service stations, meter readings, car condition photos, financial dues, penalties, and deposit refund assessments.")

add_bullet("Drop-Off Reason & Location", "Reasons: Voluntary Return, Non-payment / Default, Breakdown / Maintenance, Contract Completion. Locations: Hub Desk, Service Station, Customer Address.")
add_bullet("Meter & Photo Inspection", "Odometer Reading (KM), Odometer Photo, and 5 condition photos (Left Side, Right Side, Front Side, Back Side, Battery Photo).")
add_bullet("Financial Dues & Penalties", "Pending Dues input (₹) and Damage Penalty input (₹).")
add_bullet("Deposit Refund Assessment", "Refund status dropdown choices: Pending Assessment, Refund Approved, Deductions Applied, Forfeited.")
add_bullet("Saved Drafts & Registry Features", "Save as Draft button, Saved Drafts tab with badge count, live search bar, reason/time/city filters, 10 items per page pagination, Edit mode, Delete action, and Export to CSV.")

add_heading_2("3.4 Walk-In & Returning Partner Visit Logger Module")
add_paragraph("Manages daily hub visits, distinguishing between new candidate onboarding inquiries and returning partner operational visits.")

add_bullet("Two-Table Architecture", "Separate database storage using july_new_walkins (for candidates) and july_existing_walkins (for returning partners) with UNION listing.")
add_bullet("Entry Mode Selector", "Top selector allowing hub executives to switch between New Candidate Onboarding and Returning Partner Visit Log.")
add_bullet("Master Onboarding Phone Lookup", "Auto-populates partner name and city when typing phone number by searching approved master onboarding records.")
add_bullet("Operational Visit Tags", "Dropdown options: DM Meet, Complaint, Hisaab & Payout, Maintenance, Deposit Refund Issue, General Service, Accident, Battery Issue, Tyre Issue.")
add_bullet("Prior Visit History Modal", "Clickable history modal opening on phone lookup to view all past visits logged for a partner.")

# SECTION 4: ROLE-BASED ACCESS CONTROL (RBAC)
add_heading_1("4. Role-Based Access Control (RBAC) & User Management")
add_paragraph("The portal enforces strict role-based access control across all frontend views and backend REST API endpoints. User authorization is managed through 15 system roles mapped to a permission matrix.")

rbac_headers = ["User ID", "Username", "Role Code", "Role Name", "Generic Employee Title", "Primary Form Access & Permissions"]
rbac_data = [
    [3, "admin", "SA", "Super Admin", "System Super Admin", "Full Access: All Forms, Approvals Dashboard, User Management, Registries"],
    [41, "general_manager", "GM", "General Manager", "General Manager 1", "High-Level Approvals, Global Registries & Management Overview"],
    [17, "business", "BH", "Business Head", "Business Head 1", "Business Analytics, Global Registries & Approvals Dashboard"],
    [18, "finance_lead", "FL", "Finance Lead", "Finance Lead 1", "Financial Approvals, Hisaab & Payouts, Deposit Refund Approvals"],
    [19, "finance_executive", "FE", "Finance Executive", "Finance Executive 1", "Payment Verifications, Financial Settlement & Dues Assessment"],
    [20, "city_manager", "CM", "City Manager", "City Manager 1", "City-Level Approvals Dashboard, City Allocation & Drop-off Registries"],
    [21, "driver_manager", "DM", "Driver Manager", "Driver Manager 1", "Driver Onboarding, Driver Visits, Walk-In Logger, DM Meet Logs"],
    [23, "ops_executive", "OE", "Ops Executive", "Ops Executive 1", "Vehicle Allocation Form, Vehicle Drop-Off Form, Walk-In Logger"],
    [24, "fleet_manager", "FM", "Fleet Manager", "Fleet Manager 1", "Vehicle Allocation, Drop-Off, Maintenance & Fleet Master Logs"],
    [25, "maintenance_coordinator", "MC", "Maintenance Coordinator", "Maintenance Coordinator 1", "Vehicle Maintenance Logs, Repair Requests & Drop-off Approvals"],
    [26, "onboarding_executive", "OB", "Onboarding Executive", "Onboarding Executive 1", "Driver Onboarding Form, Walk-In Form, Vehicle Drop-Off Form"],
    [28, "support_executive", "SP", "Support Executive", "Support Executive 1", "Walk-In Logger, Support Tickets & Driver Inquiry Handling"],
    [29, "auditor", "AU", "Auditor / Compliance", "Auditor 1", "Read-only View access across all forms, registries & audit logs"],
    [31, "fleet_partner", "PT", "Fleet Partner / Operator", "Fleet Partner 1", "Operator Fleet View, Sub-driver assignments & Fleet Drop-offs"],
    [32, "driver", "DR", "Driver", "Driver 1", "Personal Profile View, Vehicle Allocation & Drop-off history"]
]

t_rbac = doc.add_table(rows=len(rbac_data)+1, cols=6)
t_rbac.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_rbac)

hdr_r = t_rbac.rows[0].cells
for idx, text in enumerate(rbac_headers):
    hdr_r[idx].text = text
    set_cell_background(hdr_r[idx], "F1F5F9")
    set_cell_padding(hdr_r[idx])
    p = hdr_r[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Inter"
        r.font.size = Pt(9.5)
        r.font.bold = True

for r_i, r_d in enumerate(rbac_data, start=1):
    row_c = t_rbac.rows[r_i].cells
    bg = "F8FAFC" if r_i % 2 == 1 else "FFFFFF"
    for c_i, val in enumerate(r_d):
        row_c[c_i].text = str(val)
        set_cell_background(row_c[c_i], bg)
        set_cell_padding(row_c[c_i])
        p = row_c[c_i].paragraphs[0]
        if c_i in [0, 2, 4]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Inter"
            r.font.size = Pt(9)
            if c_i == 1:
                r.font.bold = True

for r in t_rbac.rows:
    r.cells[0].width = Inches(0.7)
    r.cells[1].width = Inches(1.5)
    r.cells[2].width = Inches(0.8)
    r.cells[3].width = Inches(1.5)
    r.cells[4].width = Inches(1.5)
    r.cells[5].width = Inches(1.5)

# SECTION 5: AUDIT LOGGING & EVENT TRACKING ENGINE
add_heading_1("5. Audit Logging & Event Tracking Engine")
add_paragraph("The portal incorporates 6 dedicated database log tables that capture every single action, record modification, approval transition, and user authentication event in real-time.")

log_headers = ["Audit Log Table", "Target Module", "Action Types Tracked", "Audit Data Captured"]
log_data = [
    ["july_allocation_form_logs", "Allocation & Drop-Off", "CREATE, UPDATE, DELETE, DRAFT_SAVE", "Allocation ID, Executive ID, IST Timestamp, Driver ID, Vehicle No, Snapshot & JSON diff"],
    ["july_walkin_logs", "Walk-In & Partner Visits", "CREATE, UPDATE, DELETE", "Walkin ID (N1, E1), Executive ID, IST Timestamp, Remarks, Snapshot & JSON diff"],
    ["july_onboarding_logs", "Partner Onboarding", "CREATE, UPDATE, DRAFT_SAVE, RESUBMIT", "Onboarding ID, Executive ID, IST Timestamp, Full Snapshot & changed_fields JSON diff"],
    ["july_vehicle_logs", "Fleet Vehicle Master", "CREATE, UPDATE, STATUS_CHANGE", "Vehicle ID, Executive ID, IST Timestamp, Status, Full Snapshot & JSON diff"],
    ["july_approval_chain_logs", "Approvals Dashboard", "APPROVE, REJECT, REVISION, FORWARD", "Module Name, Record ID, From User ID, To User ID, Action, Approver Remarks & IST Timestamp"],
    ["july_user_login_logs", "Authentication Engine", "LOGIN, LOGOUT, FAILED_LOGIN", "User ID, Username, Full Name, Role Code, Login Time, Logout Time, IP Address & Browser User-Agent"]
]

t_log = doc.add_table(rows=len(log_data)+1, cols=4)
t_log.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_log)

hdr_l = t_log.rows[0].cells
for idx, text in enumerate(log_headers):
    hdr_l[idx].text = text
    set_cell_background(hdr_l[idx], "F1F5F9")
    set_cell_padding(hdr_l[idx])
    p = hdr_l[idx].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Inter"
        r.font.size = Pt(9.5)
        r.font.bold = True

for r_i, r_d in enumerate(log_data, start=1):
    row_c = t_log.rows[r_i].cells
    bg = "F8FAFC" if r_i % 2 == 1 else "FFFFFF"
    for c_i, val in enumerate(r_d):
        row_c[c_i].text = str(val)
        set_cell_background(row_c[c_i], bg)
        set_cell_padding(row_c[c_i])
        p = row_c[c_i].paragraphs[0]
        if c_i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = "Inter"
            r.font.size = Pt(9)
            if c_i == 0:
                r.font.bold = True

for r in t_log.rows:
    r.cells[0].width = Inches(1.8)
    r.cells[1].width = Inches(1.5)
    r.cells[2].width = Inches(1.5)
    r.cells[3].width = Inches(2.7)

# SECTION 6: DATABASE SCHEMA & ENTITY RELATIONSHIPS
add_heading_1("6. Database Schema & Entity Relationships")
add_paragraph("The database schema is structured around normalized core entity tables, application log tables, master lookup tables, and approval tracking tables.")

add_bullet("july_portal_users", "Stores executive user accounts, roles, portal_user_id (PK), employee_id, username, password_hash, role_id, city, joining_date, and account_status.")
add_bullet("july_employees", "Stores employee master records including employee_id (PK), first_name, last_name, email, phone, city, department, and joining_date.")
add_bullet("july_form_onboarding", "Master onboarding registry storing candidate personal details, KYC document URLs, address proofs, bank details, approval status, and timestamps.")
add_bullet("july_allocation_form", "Unified storage for Vehicle Allocation and Vehicle Drop-off records, distinguished by allocation_type ('Allocation' vs 'Drop-Off') and sub_type.")
add_bullet("july_new_walkins", "Stores new candidate walk-in entries including N-series auto-increment ID, candidate info, DL/Aadhaar numbers, lead channel, and joined status.")
add_bullet("july_existing_walkins", "Stores returning partner visit logs including E-series auto-increment ID, partner info, partner category, visiting reason, and visit notes.")
add_bullet("july_cities", "Master lookup table for active operating cities (Bengaluru, Mumbai, Hyderabad, Pune, Chennai).")

# SECTION 7: APPROVAL WORKFLOWS & STATE MACHINE
add_heading_1("7. Approval Workflows & State Machine")
add_paragraph("Applications submitted for onboarding or vehicle allocation undergo multi-tier state machine transitions depending on the role hierarchy.")

add_bullet("Draft State", "Initial state when an executive saves progress. Record is editable by the creator and hidden from public registries.")
add_bullet("Pending Approval State", "Submitted state where record enters the approval queue. Locked from basic edits and visible to designated approvers.")
add_bullet("Revision Required State", "Triggered when an approver sends back an application with revision comments. Re-opens form for executive corrections.")
add_bullet("Approved State", "Final active state upon successful verification. Promotes candidate to Master Registry or releases vehicle allocation.")
add_bullet("Rejected State", "Terminal state for ineligible candidates or failed verifications.")

# SECTION 8: DEPLOYMENT & DEVOPS GUIDE
add_heading_1("8. Deployment, DevOps & Operations Guide")
add_paragraph("Complete instructions for local development, production build, process management, and database pooling.")

add_heading_2("8.1 Local Development Commands")
add_bullet("Backend API Server", "Run from root directory: .venv\\Scripts\\python.exe -m uvicorn main:app --reload --host 127.0.0.1 --port 8000")
add_bullet("Frontend Dev Server", "Run from root directory: npm run dev (Launches Vite server on http://localhost:3000)")
add_bullet("Production Build Check", "Run: npm run build (Generates minified static bundle in dist/ directory)")

add_heading_2("8.2 Environment Variables Configuration")
add_paragraph("Configure the following variables in .env file:")
add_bullet("DATABASE_URL", "postgresql://postgres:postgres@localhost:5432/portal_db")
add_bullet("DB_USER / DB_PASS", "PostgreSQL database user and password credentials")
add_bullet("DB_HOST / DB_PORT", "127.0.0.1 and port 5432")

# Save Word document
doc.save(output_file)
print(f"[OK] Comprehensive 10-20 Page Portal Documentation (.docx) successfully generated at: {output_file}")
